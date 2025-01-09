import docx
import xml.etree.ElementTree as ET
import re
from collections import Counter
from pathlib import Path
import json
import logging
from datetime import datetime
from markdown_converter import MarkdownConverter

class ConRed:
    def __init__(self, config_path=None):
        """
        Initialize ConRed with optional config file
        """
        self.replacement_dict = {}
        self.replacement_counts = {'word': {}, 'xml': {}, 'markdown': {}}
        self.markdown_converter = MarkdownConverter()
        
        # Setup logging
        log_dir = Path('logs')
        log_dir.mkdir(exist_ok=True)
        logging.basicConfig(
            filename=log_dir / f'conred_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log',
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        
        if config_path:
            self.load_config(config_path)
    
    def load_config(self, config_path):
        """Load replacement dictionary from JSON config file"""
        try:
            with open(config_path, 'r') as f:
                self.replacement_dict = json.load(f)
            logging.info(f"Loaded configuration from {config_path}")
        except Exception as e:
            logging.error(f"Error loading config: {str(e)}")
            raise

    def replace_text(self, text):
        """
        Perform case-insensitive replacement while preserving replacement text case.
        Handles markdown-escaped special characters and various markdown formatting patterns.
        Returns tuple of (modified text, counter of replacements)
        """
        counts = Counter()
        modified_text = text
        
        for original, replacement in self.replacement_dict.items():
            # Create patterns for different markdown escape/formatting scenarios
            patterns = [
                # Basic word with boundaries
                r'\b' + re.escape(original) + r'\b',
                # Markdown escaped characters
                r'\b' + re.escape(original).replace('.', r'\.').replace('-', r'\-') + r'\b',
                # Inside markdown formatting (e.g., __text__, **text**)
                r'(?<=__)\s*' + re.escape(original) + r'\s*(?=__)',
                r'(?<=\*\*)\s*' + re.escape(original) + r'\s*(?=\*\*)',
                # After markdown list markers
                r'(?<=- )' + re.escape(original),
                r'(?<=\* )' + re.escape(original),
                # Inside markdown links [text]
                r'(?<=\[)' + re.escape(original) + r'(?=\])',
                # With HTML escape sequences
                r'\b' + re.escape(original).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') + r'\b'
            ]
            
            # Apply each pattern
            for pattern in patterns:
                regex = re.compile(pattern, re.IGNORECASE)
                matches = regex.findall(modified_text)
                counts[original] += len(matches)
                modified_text = regex.sub(replacement, modified_text)
        
        return modified_text, counts

    def process_word_document(self, input_path):
        """Process a Word document and count replacements while preserving formatting."""
        try:
            doc = docx.Document(input_path)
            self.replacement_counts['word'] = Counter()
            
            # Process all runs in all paragraphs to preserve formatting
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        new_text, counts = self.replace_text(run.text)
                        self.replacement_counts['word'].update(counts)
                        run.text = new_text
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text:
                                    new_text, counts = self.replace_text(run.text)
                                    self.replacement_counts['word'].update(counts)
                                    run.text = new_text
            
            # Process headers and footers
            for section in doc.sections:
                for header in [section.header, section.footer]:
                    if header:
                        for paragraph in header.paragraphs:
                            for run in paragraph.runs:
                                if run.text:
                                    new_text, counts = self.replace_text(run.text)
                                    self.replacement_counts['word'].update(counts)
                                    run.text = new_text
            
            logging.info(f"Processed Word document: {input_path}")
            return doc
            
        except Exception as e:
            logging.error(f"Error processing Word document: {str(e)}")
            raise

    def process_xml_document(self, input_path):
        """Process an XML document and count replacements."""
        try:
            # Read the file as text
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            self.replacement_counts['xml'] = Counter()
            
            # Do the replacements on the entire content
            new_content, counts = self.replace_text(content)
            self.replacement_counts['xml'].update(counts)
            
            # Return the modified content
            return new_content
            
        except Exception as e:
            logging.error(f"Error processing XML document: {str(e)}")
            raise

    def verify_counts(self):
        """Verify replacement counts match across all document types."""
        mismatches = []
        
        # Get all unique words that were replaced in any document
        all_words = set().union(
            self.replacement_counts['word'].keys(),
            self.replacement_counts['xml'].keys(),
            self.replacement_counts['markdown'].keys()
        )
        
        for word in all_words:
            word_count = self.replacement_counts['word'].get(word, 0)
            xml_count = self.replacement_counts['xml'].get(word, 0)
            md_count = self.replacement_counts['markdown'].get(word, 0)
            
            # Check for mismatches between any pair of documents
            if word_count != xml_count or word_count != md_count or xml_count != md_count:
                mismatches.append({
                    'word': word,
                    'word_doc_count': word_count,
                    'xml_doc_count': xml_count,
                    'markdown_doc_count': md_count
                })
                logging.warning(
                    f"Mismatch found for '{word}': "
                    f"Word doc: {word_count}, "
                    f"XML doc: {xml_count}, "
                    f"Markdown: {md_count}"
                )
        
        return mismatches

    def process_documents(self, word_input, xml_input, word_output, xml_output):
        """Process both documents and save results."""
        # === File Path Setup ===
        # Convert all paths to Path objects for consistent handling
        word_input = Path(word_input)
        xml_input = Path(xml_input)
        word_output = Path(word_output)
        xml_output = Path(xml_output)
        markdown_output = word_output.with_suffix('.md')  # Create markdown output path
        
        # Create output directories if they don't exist
        word_output.parent.mkdir(parents=True, exist_ok=True)
        xml_output.parent.mkdir(parents=True, exist_ok=True)
        
        # === Document Processing ===
        # Process the original Word and XML documents
        modified_doc = self.process_word_document(word_input)
        modified_xml = self.process_xml_document(xml_input)
        
        # === Markdown Conversion and Processing ===
        # Convert Word to markdown and apply replacements
        markdown_content = self.markdown_converter.convert_docx_to_markdown(word_input)
        modified_markdown, markdown_counts = self.markdown_converter.process_markdown(
            markdown_content,
            self.replacement_dict
        )
        
        # Update the counts for verification
        self.replacement_counts['markdown'] = markdown_counts
        
        # === Save Modified Documents ===
        modified_doc.save(word_output)
        with open(xml_output, 'w', encoding='utf-8') as f:
            f.write(modified_xml)
        with open(markdown_output, 'w', encoding='utf-8') as f:
            f.write(modified_markdown)
        
        # === Results Collection and Verification ===
        # Check for mismatches between document types
        mismatches = self.verify_counts()
        
        # Compile results for reporting
        results = {
            'mismatches': mismatches,
            'word_counts': dict(self.replacement_counts['word']),
            'xml_counts': dict(self.replacement_counts['xml']),
            'markdown_counts': dict(self.replacement_counts['markdown'])
        }
        
        # Log completion status
        logging.info(f"Processing complete. Mismatches found: {len(mismatches)}")
        
        return results

def main():
    # Example configuration file (config.json)
    config = {
        'Dynata': 'Panel1',
        'Toluna': 'Panel2',
        'Global Market Insite': 'Panel3',
        'M3': 'Panel4',
        'SurveyBods': 'Panel5',
        'St Modwen Homes': 'Client1',
        # Add any other sensitive terms that need redacting
    }
    
    # Save example config
    config_path = Path('config.json')
    with open(config_path, 'w') as f:
        json.dump(config, f, indent=4)
    
    # Initialize ConRed
    conred = ConRed(config_path)
    
    # Process documents
    results = conred.process_documents(
        'data/input/questionnaire.docx',
        'data/input/questionnaire.xml',
        'data/output/questionnaire_redacted.docx',
        'data/output/questionnaire_redacted.xml'
    )
    
    # Print results
    if results['mismatches']:
        print("\nWarning: Replacement count mismatches found:")
        for mismatch in results['mismatches']:
            print(f"Word '{mismatch['word']}': "
                  f"Word doc: {mismatch['word_doc_count']}, "
                  f"XML doc: {mismatch['xml_doc_count']}")
    else:
        print("\nAll replacement counts match between documents!")
    
    print("\nReplacement counts:")
    for word, count in results['word_counts'].items():
        print(f"{word}: {count} replacements")

if __name__ == "__main__":
    main()
